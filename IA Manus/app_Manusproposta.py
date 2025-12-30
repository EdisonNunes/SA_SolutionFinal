import os
import cloudconvert
import requests
import time
from docx import Document
from tkinter import Tk, filedialog, messagebox
from cloudconvert.exceptions.exceptions import UnauthorizedAccess, ConnectionError

# Configurações da API
CLOUDCONVERT_API_KEY = 'SUA_CHAVE_AQUI' # Substitua pela sua chave real

def gerar_documento_word(num_proposta, filename='temp.docx'):
    """
    Simula a geração do documento Word.
    Substitua esta lógica pela sua função real de criação do DOCX.
    """
    doc = Document()
    doc.add_heading(f'Proposta Comercial #{num_proposta}', 0)
    doc.add_paragraph('Este é um documento gerado automaticamente para conversão.')
    # Adicione aqui tabelas, logos, etc.
    doc.save(filename)
    return filename

def converter_para_pdf_api(api_key, input_file, output_filename):
    """
    Função interna para realizar a conversão via CloudConvert.
    Retorna o conteúdo binário do PDF.
    """
    cloudconvert.configure(api_key=api_key, sandbox=False)
    
    # Criar Job
    job = cloudconvert.Job.create(payload={
        "tasks": {
            "import-1": {"operation": "import/upload"},
            "convert-1": {
                "operation": "convert",
                "input": "import-1",
                "output_format": "pdf",
                "engine": "office"
            },
            "export-1": {"operation": "export/url", "input": "convert-1"}
        }
    })

    # Upload
    upload_task_id = next(task['id'] for task in job['tasks'] if task['name'] == 'import-1')
    upload_task = cloudconvert.Task.find(id=upload_task_id)
    cloudconvert.Task.upload(file_name=input_file, task=upload_task)

    # Aguardar
    job_finished = cloudconvert.Job.wait(id=job['id'])
    
    # Obter URL e baixar
    export_task = next(task for task in job_finished['tasks'] if task['name'] == 'export-1')
    file_url = export_task['result']['files'][0]['url']
    
    response = requests.get(file_url)
    response.raise_for_status()
    return response.content

def acao_gerar_pdf(num_proposta):
    """
    Função principal chamada pelo botão 'Gerar PDF'.
    """
    temp_docx = 'temp.docx'
    pdf_default_name = f"Proposta_{num_proposta}.pdf"
    
    try:
        # 1. Gerar o Word temporário
        print("Gerando documento Word temporário...")
        gerar_documento_word(num_proposta, temp_docx)

        # 2. Converter para PDF (obtendo os dados binários)
        print("Convertendo para PDF via CloudConvert...")
        pdf_content = converter_para_pdf_api(CLOUDCONVERT_API_KEY, temp_docx, pdf_default_name)

        # 3. Abrir diálogo para salvar arquivo
        root = Tk()
        root.withdraw() # Esconde a janela principal do tkinter
        root.attributes("-topmost", True) # Garante que o diálogo apareça na frente
        
        caminho_salvamento = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            initialfile=pdf_default_name,
            filetypes=[("Arquivos PDF", "*.pdf")],
            title="Escolha onde salvar sua Proposta"
        )
        root.destroy()

        if caminho_salvamento:
            # 4. Gravar o arquivo no local escolhido
            with open(caminho_salvamento, 'wb') as f:
                f.write(pdf_content)
            messagebox.showinfo("Sucesso", f"PDF gerado e salvo com sucesso em:\n{caminho_salvamento}")
        else:
            print("Operação de salvamento cancelada pelo usuário.")

    except UnauthorizedAccess:
        messagebox.showerror("Erro de API", "Chave do CloudConvert inválida ou e-mail não verificado.")
    except Exception as e:
        messagebox.showerror("Erro Inesperado", f"Ocorreu um erro: {str(e)}")
    
    finally:
        # 5. Apagar o arquivo temporário
        if os.path.exists(temp_docx):
            os.remove(temp_docx)
            print("Arquivo temporário temp.docx removido.")

if __name__ == "__main__":
    # Exemplo de execução:
    # No seu sistema, este seria o comando disparado pelo botão da sua interface
    num_exemplo = "12345"
    acao_gerar_pdf(num_exemplo)
