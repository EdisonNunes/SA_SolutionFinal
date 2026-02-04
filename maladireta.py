from docx import Document   # pip install python-docx
from docx.shared import Pt
from datetime import datetime
# from converte import converter_docx_para_pdf
import streamlit as st
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

CLOUDCONVERT_API_KEY = st.secrets["cloudconvert"]["CLOUDCONVERT_API_KEY"]

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Função auxiliar para aplicar bordas em uma célula.
    Uso: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "4F81BD"}, ...)
    sz: largura em oitavos de ponto (12 = 1.5pt)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'start', 'bottom', 'end'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, value in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(value))


def listar_alinhamentos_da_tabela(docx_path):
    doc = Document(docx_path)
    
    # Dicionário para traduzir os códigos da biblioteca
    traducao = {
        WD_ALIGN_PARAGRAPH.LEFT: "Esquerda",
        WD_ALIGN_PARAGRAPH.CENTER: "Centralizado",
        WD_ALIGN_PARAGRAPH.RIGHT: "Direita",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "Justificado",
        None: "Padrão (Esquerda)"
    }

    for t_idx, tabela in enumerate(doc.tables):
        # print(f"\nTabela {t_idx + 1}:")
        for r_idx, linha in enumerate(tabela.rows):
            for c_idx, celula in enumerate(linha.cells):
                # Acessamos o alinhamento do parágrafo dentro da célula
                alinhamento = celula.paragraphs[0].alignment
                # print(f"  [L{r_idx}, C{c_idx}]: {traducao.get(alinhamento)}")



# ======================================================
# 1. Função auxiliar: substituição de TAGS
# ======================================================
def substituir_tags(paragrafo, tags: dict):
    """
    Substitui tags dentro de um parágrafo do Word.
    Se houver campos (ex: numeração de página), faz substituição run-a-run para não quebrá-los.
    Caso contrário, substitui no texto do parágrafo (pode resetar formatação interna).
    """
    
    # 1. Verificar se existem campos (Field Characters) no parágrafo
    tem_campo = False
    for run in paragrafo.runs:
        if "w:fldChar" in run._element.xml:
            tem_campo = True
            break
            
    if tem_campo:
        # Modo Seguro: Substituir apenas dentro de cada Run para preservar os campos
        for run in paragrafo.runs:
            if run.text:
                texto_run = run.text
                mudou = False
                for chave, valor in tags.items():
                    if chave in texto_run:
                        texto_run = texto_run.replace(chave, str(valor))
                        mudou = True
                if mudou:
                    run.text = texto_run
    else:
        # Modo Padrão: Substituir no texto completo (garante que tags quebradas sejam achadas)
        if not paragrafo.text:
            return
        
        texto_atual = paragrafo.text
        # Otimização: só atribui se houver mudança
        for chave, valor in tags.items():
            if chave in texto_atual:
                paragrafo.text = paragrafo.text.replace(chave, str(valor))

# ======================================================
# 2. Geração do documento Word
# ======================================================
def gerar_documento_word(
    supabase,
    id_proposta: int,
    caminho_template: str,
    caminho_saida: str
):
    # -------------------------------
    # Buscar dados da proposta
    # -------------------------------
    # Exemplo de uso
    # listar_alinhamentos_da_tabela(caminho_template)


    proposta = (
        supabase.table("propostas")
        .select("*")
        .eq("id_proposta", id_proposta)
        .single()
        .execute()
    ).data

    cliente = (
        supabase.table("clientes")
        .select("*")
        .eq("id", proposta["id_cliente"])
        .single()
        .execute()
    ).data

    itens = (
        supabase.table("itens_proposta")
        .select("*")
        .eq("id_proposta", id_proposta)
        .execute()
    ).data

    # -------------------------------
    # Carregar template
    # -------------------------------
    doc = Document(caminho_template)

    # -------------------------------
    # Criar dicionário de TAGS
    # -------------------------------
    data_emissao_formatada = datetime.strptime(
        proposta["data_emissao"], "%Y-%m-%d"
    ).strftime("%d/%m/%Y")

    tags = {
        "{{NUM_PROPOSTA}}": proposta["num_proposta"],
        "{{DATA_EMISSAO}}": data_emissao_formatada,
        "{{VALIDADE}}": proposta["validade"],
        "{{COND_PAGAMENTO}}": proposta["cond_pagamento"],
        "{{REFERENCIA}}": proposta["referencia"],

        # Dados do cliente
        "{{ID}}": cliente["id"],
        "{{CLIENTE}}": cliente["empresa"],
        "{{CNPJ}}": cliente["cnpj"],
        "{{ENDERECO}}": cliente["endereco"],
        "{{CIDADE}}": cliente["cidade"],
        "{{UF}}": cliente["uf"],
        "{{CONTATO}}": cliente["contato"],
        "{{EMAIL}}": cliente["email"],
        "{{TELEFONE}}": cliente["telefone"],
    }

    # -------------------------------
    # Substituir tags no documento
    # -------------------------------

    for p in doc.paragraphs:
        substituir_tags(p, tags)
        # Arial 8 somente para conteúdo
        for run in p.runs:
            # print('run: ',run.font.name, run.font.size, run.element.text) 
            run.font.name = "Arial"
            run.font.size = Pt(8)

    # -------------------------------
    # Substituir tags em todas as tabelas
    # -------------------------------
    for tabela in doc.tables:
        #print('tabela: ',tabela.rows[0].cells)
        for linha in tabela.rows:
            for cel in linha.cells: 
                for p in cel.paragraphs:
                    substituir_tags(p, tags)
                    # Arial 8
                    for run in p.runs:
                        # print('run: ',run.font.name, run.font.size, cel.paragraphs[0].alignment)
                        run.font.name = "Arial"
                        run.font.size = Pt(8)
    # -------------------------------
    # Substituir tags no rodapé
    # -------------------------------
    for section in doc.sections:
        footer = section.footer
        # print('footer.is_linked_to_previous =',footer.is_linked_to_previous)
        for p in footer.paragraphs:
            substituir_tags(p, tags)
            for run in p.runs:
                # print('run: ',run.font.name, run.font.size, run.element.text)
                run.font.name = "Arial"
                run.font.size = Pt(5) # Tamanho 5

    # ================================
    # 4) TABELA DE ITENS (SEGUNDA TABELA!)
    # ================================
    tabela_itens = doc.tables[1]   # <-- AQUI!

    # Validar número de colunas
    if len(tabela_itens.rows[0].cells) != 8:
        raise ValueError("A tabela do template deve ter 8 colunas.")

    # Remover TODAS as linhas após o cabeçalho
    while len(tabela_itens.rows) > 1:
        tabela_itens._element.remove(tabela_itens.rows[-1]._element)


    # ALINHAMENTO DO CABEÇALHO (Linha 0)
    for celula in tabela_itens.rows[0].cells:
        celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for par in celula.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # ================================
    # 5) ADICIONAR ITENS
    # ================================
    total_proposta = 0

    for idx, item in enumerate(itens, start=1):
        qtd = float(item["qtd"])
        preco = float(item["preco_unitario"])
        desconto = float(item["desconto"])

        total_item = (qtd * preco) - (preco * desconto / 100)
        total_proposta += total_item

        linha = tabela_itens.add_row().cells

        linha[0].text = str(idx)  
        
        linha[1].text = item["codigo_servico"]
        linha[2].text = item["descricao_servico"]
        linha[3].text = item["prazo_ddl"]
        linha[4].text = str(int(qtd))
        linha[5].text = f"R$ {preco:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        linha[6].text = f"{desconto:.2f}%"
        linha[7].text = f"R$ {total_item:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        # aplica fonte Arial 8
        for cel in linha:
            for par in cel.paragraphs:
                for run in par.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)


    # --- APLICAÇÃO DE ALINHAMENTOS NAS NOVAS LINHAS ---
        for i, cel in enumerate(linha):
            # Regra 3: Alinhamento Vertical CENTRO para toda a tabela
            cel.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Regra 2: Alinhamento Horizontal por coluna
            if i in [0, 1, 3, 4, 6]: # Colunas 0, 1, 3, 4 (e 6 por ser Qtd) -> Centralizado
                cel.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i == 2:          # Coluna 2 -> Esquerda
                cel.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif i in [5, 7]:  # Colunas 5, 7 -> Direita
                cel.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


    # Configuração da borda (1.5pt, Cor 4F81BD)
    borda_config = {
        "sz": 12, 
        "val": "single", 
        "color": "4F81BD"
    }

    # ================================
    # 6) LINHA EM BRANCO
    # ================================
    linha_vazia = tabela_itens.add_row().cells
    borda_nula = {"val": "nil"}
    
    # Define kwargs para borda nula em todos os lados
    kwargs_borda_nula = {
        "top": borda_nula, 
        "bottom": borda_nula, 
        "start": borda_nula, 
        "end": borda_nula
    }

    for cel in linha_vazia:
        par = cel.paragraphs[0]
        run = par.add_run(" ")
        run.font.name = "Arial"
        run.font.size = Pt(8)
        
        # Borda nula em todas as colunas
        set_cell_border(cel, **kwargs_borda_nula)

    # ================================
    # 7) LINHA SUB.TOTAL
    # ================================
    linha_total = tabela_itens.add_row().cells

    linha_total[5].text = "Sub.Total:"
    linha_total[7].text = (
        f"R$ {total_proposta:,.2f}"
        .replace(",", "X").replace(".", ",").replace("X", ".")
    )

    for cel in linha_total:
        for par in cel.paragraphs:
            for run in par.runs:
                run.font.name = "Arial"
                run.font.size = Pt(8)
                run.font.bold = True

    # Aplique o mesmo padrão para as linhas de total:
    for i, cel in enumerate(linha_total):
        cel.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Borda nula em todas as colunas (base)
        set_cell_border(cel, **kwargs_borda_nula)

        # Aplicar APENAS a borda SUPERIOR nas colunas 6 e 7
        if i in [6, 7]:
            set_cell_border(
                cel, 
                top=borda_config
            )
            
        # Alinhamento para os campos de total
        cel.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


    doc.save('temp.docx')
    # listar_alinhamentos_da_tabela('temp.docx')
    # return caminho_saida
    return 'temp.docx', proposta["num_proposta"]
